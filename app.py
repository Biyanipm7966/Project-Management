from flask import Flask, render_template, request, send_file, jsonify
import dash
from dash import dcc, html, Input, Output, State
import dash_bootstrap_components as dbc
import pandas as pd
import plotly.express as px
from datetime import datetime
from docx import Document
import io

# Initialize Flask app
app = Flask(__name__)

# Sample project data
dashboard_data = {
    'project_name': 'Punch Card',
    'project_logo': 'https://example.com/logo.png',
    'status': 'On Track',
    'milestones': [
        {'name': 'Kickoff', 'date_from': '2024-09-24', 'date_to': '2024-09-24', 'status': 'Completed'},
        {'name': 'Brainstorming', 'date_from': '2024-09-25', 'date_to': '2024-10-08', 'status': 'Completed'},
        {'name': 'Development', 'date_from': '2024-10-09', 'date_to': '2024-10-29', 'status': 'In-Progress'},
        {'name': 'Testing and Quality Assurance', 'date_from': '2024-10-30', 'date_to': '2024-11-12', 'status': 'Scheduled'},
        {'name': 'Finalizing', 'date_from': '2024-11-13', 'date_to': '2024-11-26', 'status': 'Scheduled'},
        {'name': 'Release', 'date_from': '2024-11-27', 'date_to': '2024-11-27', 'status': 'Scheduled'},
    ]
}

# Convert milestones to a pandas DataFrame
df = pd.DataFrame(dashboard_data['milestones'])

# Initialize Dash app
dash_app = dash.Dash(
    __name__,
    server=app,
    url_base_pathname='/dashboard/',
    external_stylesheets=[dbc.themes.BOOTSTRAP]
)

# Dash App layout
dash_app.layout = html.Div([
    dbc.Form([
        dbc.Row([
            dbc.Col(dbc.Input(id='task-name', type='text', placeholder='Task Name', required=True), width=4),
            dbc.Col(dbc.Input(id='start-date', type='date', placeholder='Start Date', required=True), width=2),
            dbc.Col(dbc.Input(id='end-date', type='date', placeholder='End Date', required=True), width=2),
            dbc.Col(dbc.Button("Add Task", id='add-task-btn', color='primary', n_clicks=0), width=2)
        ], className='mb-3'),
    ]),
    dcc.Graph(id='gantt-chart')
])

# Callback to add new tasks to the DataFrame and update Gantt Chart
@dash_app.callback(
    Output('gantt-chart', 'figure'),
    Input('add-task-btn', 'n_clicks'),
    [State('task-name', 'value'),
     State('start-date', 'value'),
     State('end-date', 'value')]
)
def update_gantt_chart(n_clicks, task_name, start_date, end_date):
    global df
    if n_clicks > 0 and task_name and start_date and end_date:
        try:
            start_date = datetime.strptime(start_date, '%Y-%m-%d').date()
            end_date = datetime.strptime(end_date, '%Y-%m-%d').date()
            if end_date >= start_date:
                new_row = {'name': task_name, 'date_from': start_date, 'date_to': end_date, 'status': 'Pending'}
                df = pd.concat([df, pd.DataFrame([new_row])], ignore_index=True)
        except ValueError:
            pass

    df['date_from'] = pd.to_datetime(df['date_from'])
    df['date_to'] = pd.to_datetime(df['date_to'])
    df['date_to'] = df.apply(
        lambda row: row['date_to'] if row['date_to'] > row['date_from'] else row['date_from'] + pd.Timedelta(days=1),
        axis=1
    )
    df_sorted = df.sort_values(by='date_from', ascending=True)

    fig = px.timeline(df_sorted, x_start="date_from", x_end="date_to", y="name", title="Project Gantt Chart")
    fig.update_yaxes(categoryorder='array', categoryarray=df_sorted['name'].tolist()[::-1])
    fig.update_layout(xaxis_title="Date", yaxis_title="Tasks")

    return fig

@app.route('/generate-status-report', methods=['POST'])
def generate_status_report():
    try:
        data = request.get_json()
        name = data.get('name', 'Untitled Project')
        date = data.get('date', 'No Date Provided')
        status = data.get('status', 'No Status Provided')
        phase = data.get('phase', 'No Phase Provided')
        startDate = data.get('startDate', 'No Start Date Provided')
        endDate = data.get('endDate', 'No End Date Provided')
        sDescription = data.get('sDescription', ' ')
        accomplishments = data.get('accomplishments', 'No Accomplishments Provided')
        plannedwork = data.get('plannedwork', 'No Planned Work Provided')
        currentissues = data.get('currentissues', 'No Current Issues Provided')
        risks = data.get('risks', 'No Risks Provided')
        budget = data.get('budget', 'No Budget Provided')
        team = data.get('team', 'No Team Provided')
        decisions = data.get('decisions', 'No Decisions Provided')

        # Create Word document
        doc = Document()
        doc.add_heading(f"Status Report", level=1)
        doc.add_paragraph(f"Project Name: {name}")
        doc.add_paragraph(f"Report Date: {date}")
        doc.add_paragraph(f"Project Status: {status}")
        doc.add_paragraph(f"Project Phase: {phase}")
        doc.add_paragraph(f"Project Start Date: {startDate}")
        doc.add_paragraph(f"Project End Date: {endDate}")
        doc.add_paragraph(f"Key Milestones: {sDescription}")
        doc.add_paragraph(f"Accomplishments: {accomplishments}")
        doc.add_paragraph(f"Planned Work: {plannedwork}")
        doc.add_paragraph(f"Current Issues: {currentissues}")
        doc.add_paragraph(f"Risks: {risks}")
        doc.add_paragraph(f"Budget: {budget}")
        doc.add_paragraph(f"Team: {team}")
        doc.add_paragraph(f"Decisions Needed: {decisions}")

        # Save the document to a BytesIO buffer
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        # Return the Word document as a downloadable file
        return send_file(
            buffer,
            as_attachment=True,
            download_name=f"{date}_Status_Report.docx",
            mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/generate-progress-report', methods=['POST'])
def generate_progress_report():
    try:
        data = request.get_json()
        pname = data.get('pname', 'Untitled Project')
        rpauth = data.get('rpauth', 'No Author Provided')
        rpdate = data.get('rpdate', 'No Date Provided')
        objective = data.get('objective', 'No Objective Provided')
        tasks = data.get('tasks', 'No Completed Tasks Provided')
        milestones = data.get('milestones', 'No Reached Milestones Provided')
        obstacles = data.get('obstacles', 'No Encountered Obstacles Provided')
        delays = data.get('delays', 'No Delays Provided')
        plannedtasks = data.get('plannedtasks', 'No Planned Tasks Provided')
        deadlines = data.get('deadlines', 'No Deadlines Provided')
        metrices = data.get('metrices', 'No Metrices and KPI Provided')
        action = data.get('action', 'No Action Items Provided')
        notes = data.get('notes', 'No Notes Items Provided')

        # Create Word document
        doc = Document()
        doc.add_heading(f"Progress Report", level=1)
        doc.add_paragraph(f"Project Name: {pname}")
        doc.add_paragraph(f"Report Author: {rpauth}")
        doc.add_paragraph(f"Report Date: {rpdate}")
        doc.add_paragraph(f"Objectives for Reporting Period: {objective}")
        doc.add_paragraph(f"Tasks Completed: {tasks}")
        doc.add_paragraph(f"Milesones Reached: {milestones}")
        doc.add_paragraph(f"Obstacles Encountered: {obstacles}")
        doc.add_paragraph(f"Delays: {delays}")
        doc.add_paragraph(f"Planned Tasks: {plannedtasks}")
        doc.add_paragraph(f"Deadlines: {deadlines}")
        doc.add_paragraph(f"Metrices and KPI: {metrices}")
        doc.add_paragraph(f"Action Items: {action}")
        doc.add_paragraph(f"Aditional Notes: {notes}")

        # Save the document to a BytesIO buffer
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        # Return the Word document as a downloadable file
        return send_file(
            buffer,
            as_attachment=True,
            download_name=f"{rpdate}_Progress_Report.docx",
            mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/index.html')
def home():
    return render_template('index.html')

@app.route('/charter.html')
def dashboard():
    return render_template('charter.html', data=dashboard_data)

@app.route('/milestones.html')
def milestones():
    return render_template('milestones.html')

@app.route('/budget.html')
def budget():
    return render_template('budget.html')

@app.route('/wbs.html')
def wbs():
    return render_template('wbs.html')

@app.route('/aoa.html')
def aoa():
    return render_template('aoa.html')

@app.route('/sow.html')
def sow():
    return render_template('sow.html')

@app.route('/risks.html')
def risks():
    return render_template('risks.html')

@app.route('/reports.html')
def reports():
    return render_template('reports.html')

@app.route('/team.html')
def team():
    return render_template('team.html')

@app.route('/review.html')
def review():
    return render_template('review.html')

if __name__ == '__main__':
    app.run(debug=True)
